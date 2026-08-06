"""RiskPilotDocumentParser 真实格式解析测试。"""

from __future__ import annotations

import hashlib
import io
from collections.abc import Callable

import fitz
import pytest
from docx import Document as DocxDocument

from domain import DocumentParserPort, DocumentVersion
from domain.errors import InvalidDocumentContent, UnsupportedDocumentType
from infra.document_processing import RiskPilotDocumentParser


def _version(
    content: bytes,
    *,
    mime_type: str,
    version_id: str = "ver_001",
) -> DocumentVersion:
    return DocumentVersion(
        version_id=version_id,
        document_id="doc_001",
        version_number=1,
        object_key="objects/source",
        sha256=hashlib.sha256(content).hexdigest(),
        mime_type=mime_type,
        size_bytes=len(content),
        created_at=100.0,
    )


def _parser() -> RiskPilotDocumentParser:
    return RiskPilotDocumentParser(
        clock=lambda: 123.0,
        id_factory=lambda: "parse_001",
    )


def _pdf_bytes(*page_writers: Callable[[fitz.Page], None]) -> bytes:
    document = fitz.open()
    for writer in page_writers:
        page = document.new_page()
        writer(page)
    content = document.tobytes()
    document.close()
    return content


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("数据出境合同", level=1)
    document.add_paragraph("境外接收方应承担安全保护责任。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "地区"
    table.cell(1, 1).text = "EU"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestParserContract:
    def test_satisfies_port(self) -> None:
        assert isinstance(_parser(), DocumentParserPort)


class TestPdfParsing:
    def test_preserves_real_page_numbers(self) -> None:
        content = _pdf_bytes(
            lambda page: page.insert_text((72, 72), "page one"),
            lambda page: page.insert_text((72, 72), "page two"),
        )
        snapshot = _parser().parse(
            _version(content, mime_type="application/pdf"),
            content,
        )
        assert snapshot.page_count == 2
        assert [page.page_number for page in snapshot.pages] == [1, 2]
        assert snapshot.pages[0].text == "page one"
        assert snapshot.pages[1].text == "page two"
        assert snapshot.warnings == []

    def test_blank_page_routes_to_ocr(self) -> None:
        content = _pdf_bytes(
            lambda page: page.insert_text((72, 72), "native"),
            lambda page: None,
        )
        snapshot = _parser().parse(
            _version(content, mime_type="application/pdf"),
            content,
        )
        assert snapshot.pages[1].extraction_method == "empty"
        assert "需要 OCR" in snapshot.pages[1].warnings[0]
        assert any("第 2 页" in warning for warning in snapshot.warnings)

    def test_broken_pdf_rejected(self) -> None:
        content = b"%PDF-broken"
        with pytest.raises(InvalidDocumentContent, match="PDF"):
            _parser().parse(
                _version(content, mime_type="application/pdf"),
                content,
            )


class TestDocxAndTextParsing:
    def test_docx_extracts_paragraphs_and_tables(self) -> None:
        content = _docx_bytes()
        snapshot = _parser().parse(
            _version(
                content,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            content,
        )
        assert snapshot.page_count == 1
        assert "数据出境合同" in snapshot.pages[0].text
        assert "境外接收方" in snapshot.pages[0].text
        assert snapshot.pages[0].tables[0].row_count == 2
        assert "| 字段 | 值 |" in snapshot.pages[0].tables[0].markdown
        assert "逻辑单页" in snapshot.warnings[0]

    @pytest.mark.parametrize(
        "mime_type",
        ["text/plain", "text/markdown"],
    )
    def test_text_formats_use_logical_single_page(self, mime_type: str) -> None:
        content = "第一条 数据出境".encode()
        snapshot = _parser().parse(
            _version(content, mime_type=mime_type),
            content,
        )
        assert snapshot.page_count == 1
        assert snapshot.pages[0].text == "第一条 数据出境"
        assert "逻辑单页" in snapshot.warnings[0]

    def test_unsupported_mime_rejected(self) -> None:
        content = b"binary"
        with pytest.raises(UnsupportedDocumentType):
            _parser().parse(
                _version(content, mime_type="application/octet-stream"),
                content,
            )
