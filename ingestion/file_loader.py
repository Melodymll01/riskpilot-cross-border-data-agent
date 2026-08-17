"""文件加载器：支持 PDF / TXT / DOCX 读取，输出统一的 RawDocument。"""

import logging
from pathlib import Path

from docx import Document as DocxDocument

from ingestion.pdf_extractor import PDFExtractor
from ingestion.unified_loader import RawDocument

logger = logging.getLogger(__name__)


class FileLoader:
    """读取本地文件并返回统一文档对象。

    PDF 解析委托给 `PDFExtractor`，本类只负责按后缀分发，符合单一职责。
    可通过依赖注入传入自定义配置的 `PDFExtractor`，便于按数据源切换策略。
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

    def __init__(self, pdf_extractor: PDFExtractor | None = None):
        self.pdf_extractor = pdf_extractor or PDFExtractor()

    def load(self, file_path: str, original_filename: str | None = None) -> RawDocument:
        """
        加载指定路径的文件，返回 RawDocument。

        Args:
            file_path: 文件在服务器上的实际路径
            original_filename: 用户上传时的原始文件名
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        name = original_filename or path.name

        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {suffix}，仅支持 {self.SUPPORTED_EXTENSIONS}")

        logger.info(f"正在加载文件: {name} ({suffix})")

        if suffix == ".pdf":
            content = self._load_pdf(path)
        elif suffix == ".txt":
            content = self._load_txt(path)
        elif suffix == ".docx":
            content = self._load_docx(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

        # 提取标题：使用无后缀的文件名
        title = Path(name).stem

        return RawDocument(
            content=content,
            source_type="file",
            source_name=name,
            title=title,
            source_url=None,
        )

    def _load_pdf(self, path: Path) -> str:
        """读取 PDF 文件全部文本，委托给 PDFExtractor。"""
        return self.pdf_extractor.extract(path)

    def _load_txt(self, path: Path) -> str:
        """读取纯文本文件，自动检测编码。"""
        for encoding in ("utf-8", "gbk", "gb2312", "utf-16", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue
        # 最终兜底
        return path.read_text(encoding="utf-8", errors="replace")

    def _load_docx(self, path: Path) -> str:
        """读取 DOCX 文件全部段落和表格文本。

        python-docx 的 doc.paragraphs 只返回正文段落，不包含表格。
        法规文档中大量表格（评估指标、分类标准等）必须单独提取。
        通过遍历 document.element.body 按原始顺序交替提取段落和表格。
        """
        doc = DocxDocument(str(path))
        parts = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # 去掉命名空间前缀

            if tag == "p":
                # 段落
                text = element.text
                # element.text 只取直接子节点文本，需要递归取全部
                text = "".join(node.text or "" for node in element.iter() if node.text)
                text = text.strip()
                if text:
                    parts.append(text)

            elif tag == "tbl":
                # 表格：按行提取，每行用 " | " 分隔单元格
                table_rows = []
                for tr in element.iter(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"
                ):
                    cells = []
                    for tc in tr.iter(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"
                    ):
                        cell_text = "".join(
                            node.text or "" for node in tc.iter() if node.text
                        ).strip()
                        cells.append(cell_text)
                    if any(cells):
                        table_rows.append(" | ".join(cells))
                if table_rows:
                    parts.append("\n".join(table_rows))

        return "\n\n".join(parts)
