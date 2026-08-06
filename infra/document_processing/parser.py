"""原始文件字节到页级解析快照的默认实现。"""

from __future__ import annotations

import io
import uuid
from collections.abc import Callable

import fitz
from docx import Document as DocxDocument

from domain.document_content import DocumentParseSnapshot, ParsedPage, ParsedTable
from domain.documents import DocumentVersion
from domain.errors import InvalidDocumentContent, UnsupportedDocumentType

PARSER_NAME = "riskpilot-document-parser"
PARSER_VERSION = "1.0.0"


class RiskPilotDocumentParser:
    """PDF 保留真实页码，Office/文本格式先输出明确标注的逻辑单页。"""

    def __init__(
        self,
        *,
        clock: Callable[[], float],
        id_factory: Callable[[], str] | None = None,
    ) -> None:
        self._clock = clock
        self._id_factory = id_factory or (lambda: f"parse_{uuid.uuid4().hex[:16]}")

    def parse(
        self,
        version: DocumentVersion,
        content: bytes,
    ) -> DocumentParseSnapshot:
        if version.mime_type == "application/pdf":
            pages, warnings = self._parse_pdf(content)
        elif version.mime_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            pages, warnings = self._parse_docx(version, content)
        elif version.mime_type in {"text/plain", "text/markdown"}:
            pages, warnings = self._parse_text(content)
        else:
            raise UnsupportedDocumentType(f"解析器不支持 MIME 类型 {version.mime_type!r}")
        return DocumentParseSnapshot(
            snapshot_id=self._id_factory(),
            document_version_id=version.version_id,
            parser_name=PARSER_NAME,
            parser_version=PARSER_VERSION,
            source_sha256=version.sha256,
            pages=pages,
            warnings=warnings,
            parsed_at=self._clock(),
        )

    @staticmethod
    def _parse_pdf(content: bytes) -> tuple[list[ParsedPage], list[str]]:
        try:
            document = fitz.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise InvalidDocumentContent("PDF 文件无法打开") from exc
        if document.page_count < 1:
            document.close()
            raise InvalidDocumentContent("PDF 不包含页面")

        pages: list[ParsedPage] = []
        warnings: list[str] = []
        try:
            for page_index in range(document.page_count):
                page_number = page_index + 1
                page = document.load_page(page_index)
                try:
                    text = (page.get_text("text") or "").strip()
                except Exception as exc:
                    text = ""
                    warnings.append(f"第 {page_number} 页文本层抽取失败: {type(exc).__name__}")
                if text:
                    pages.append(
                        ParsedPage(
                            page_number=page_number,
                            text=text,
                            extraction_method="native",
                        )
                    )
                else:
                    page_warning = "未检测到文本层，需要 OCR"
                    warnings.append(f"第 {page_number} 页{page_warning}")
                    pages.append(
                        ParsedPage(
                            page_number=page_number,
                            extraction_method="empty",
                            warnings=[page_warning],
                        )
                    )
        finally:
            document.close()
        return pages, warnings

    @staticmethod
    def _parse_docx(
        version: DocumentVersion,
        content: bytes,
    ) -> tuple[list[ParsedPage], list[str]]:
        try:
            document = DocxDocument(io.BytesIO(content))
        except Exception as exc:
            raise InvalidDocumentContent("DOCX 文件无法解析") from exc

        paragraphs = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        tables: list[ParsedTable] = []
        for table_index, table in enumerate(document.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ").replace("|", "\\|") for cell in row.cells
                ]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            normalized = [row + [""] * (column_count - len(row)) for row in rows]
            header = normalized[0]
            markdown_lines = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(["---"] * column_count) + " |",
            ]
            markdown_lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
            tables.append(
                ParsedTable(
                    table_id=f"{version.version_id}:p1:t{table_index}",
                    page_number=1,
                    markdown="\n".join(markdown_lines),
                    row_count=len(normalized),
                    column_count=column_count,
                )
            )

        text = "\n\n".join(paragraphs)
        if not text and not tables:
            raise InvalidDocumentContent("DOCX 未提取到正文或表格")
        warning = "DOCX 当前按逻辑单页解析，不提供物理页码"
        return (
            [
                ParsedPage(
                    page_number=1,
                    text=text,
                    extraction_method="native",
                    tables=tables,
                    warnings=[warning],
                )
            ],
            [warning],
        )

    @staticmethod
    def _parse_text(content: bytes) -> tuple[list[ParsedPage], list[str]]:
        try:
            text = content.decode("utf-8-sig").strip()
        except UnicodeDecodeError as exc:
            raise InvalidDocumentContent("文本文件必须使用 UTF-8 编码") from exc
        if not text:
            raise InvalidDocumentContent("文本文件未包含有效内容")
        warning = "文本文件按逻辑单页解析"
        return (
            [
                ParsedPage(
                    page_number=1,
                    text=text,
                    extraction_method="native",
                    warnings=[warning],
                )
            ],
            [warning],
        )
